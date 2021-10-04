import os
import random
from docx import Document


def generate_name():
    first = ("Super", "Retarded", "Great", "Sexy", "Vegan", "Brave", "Shy", "Cool", "Poor", "Rich", "Fast", "Gummy",
             "Yummy", "Masked", "Unusual", "American", "Bisexual", "MLG", "Mlg", "lil", "Lil")

    return random.choice(first)


def generate_surname():
    second = ("Coder", "Vegan", "Man", "Hacker", "Horse", "Bear", "Goat", "Goblin", "Learner", "Killer", "Woman",
              "Programmer", "Spy", "Stalker", "Spooderman", "Carrot", "Goat", "Quickscoper", "Quickscoper")

    return random.choice(second)


def generate_date():
    possible_month = ('1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12')
    possible_day = ('1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15',
                    '16', '17', '18', '19', '20', '21', '22', '23', '24', '25', '26', '27', '28', '29', '30', '31')

    day = random.choice(possible_day)
    month = random.choice(possible_month)

    if (month in [4, 6, 9, 11] and day == 31) or (month == 2 and day in [29, 30, 31]):
        return generate_date()
    else:
        return str(month)+'/'+str(day)+'/2019'


def generate_id():
    return str(random.randint(1000000000000, 9999999999999))


def generate_telephone():
    return '0' + str(random.randint(700000000, 799999999))


def generate_email(first, second):
    possible_email = ('gmail', 'hotmail', 'yahoo')
    possible_extension = ('.com', '.ro', '.en')
    return first + '.' + second + '@' + random.choice(possible_email) + random.choice(possible_extension)


def generate_selection(wished_courses):
    length = random.randint(1, len(wished_courses) + 1)
    the_string = ''
    counter = 0

    for counter in range(length):
        option = random.choice(wished_courses)
        if option not in the_string:
            the_string += option + ', '
            counter +=1

    the_string = the_string[:-2]

    return the_string


def generate_document(file_name):
    document = Document()
    document.add_heading('Registration Sheet')
    document.add_paragraph('Registration date:  ' + generate_date())
    name = generate_name()
    document.add_paragraph('Name: ' + name)
    surname = generate_surname()
    document.add_paragraph('Surname: ' + surname)
    document.add_paragraph('ID: ' + generate_id())
    document.add_paragraph('Telephone: ' + generate_telephone())
    document.add_paragraph('Email: ' + generate_email(name, surname))
    document.add_paragraph('Available courses: ')
    document.add_paragraph('PB1 -  Programming Basic 1')
    document.add_paragraph('PB2 -  Programming Basic 2')
    document.add_paragraph('PB3 -  Programming Basic 3')
    document.add_paragraph('TD1 -  Technical Drawing 1')
    document.add_paragraph('TD2 -  Technical Drawing 2')
    document.add_paragraph('CD1 -  Contemporary Dance 1')
    document.add_paragraph('CD2 -  Contemporary Dance 2')
    document.add_paragraph('CB1 -  Cooking Basic 1')
    document.add_paragraph('AE1 -  Advanced Electrician 1')
    document.add_paragraph('EB1 -  Electrician Basic 1')
    document.add_paragraph('Desired courses(The list of courses in order of priority): '
                           + generate_selection(('PB1', 'PB2', 'PB3', 'TD1', 'TD2', 'CD1', 'CD2', 'CB1', 'EB1', 'AE1')))
    document.save(os.path.join(os.getcwd(), 'applicants', 'enrollment' + file_name + '.docx'))


if __name__ == "__main__":
    # we use main for local testing of module

    for i in range(20):
        generate_document(str(i))
