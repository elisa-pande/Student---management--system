import datetime
import json
import os
from docx import Document

# lambda function for eliminating space at the beginning of string
remove_starting_space = lambda param: param[1:] if param[0] == " " else param


def modify_registration_form(courses_details: str, folder_path: str, debug: bool = False) -> bool:
    """
    Function that modifies the admission template and creates the new model
    :param courses_details: text to add which details the courses abbreviation and name
    :param folder_path: folder location
    :param debug:
    :return: True if everything was ok and model created
    """
    file = os.path.join(os.getcwd(), folder_path, "template_registration.txt")

    if debug:
        print("DEBUG: Text for modifying admission form:\n", courses_details, '\n', '_' * 20)
        print("DEBUG: Folder path for location of admission: ", file, '\n', '_' * 20)

    # get the data from the template and read it
    text = open(file, 'r').read()
    text = text.replace('XXXXX', courses_details)
    # write the information in the template docx file
    document = Document()
    document.add_heading(text.split("\n")[0])
    document.add_paragraph(text[15:])
    document.save(os.path.join(os.getcwd(), folder_path, 'template_registration.docx'))
    # check if the docs file was created
    if os.path.exists(os.path.join(os.getcwd(), folder_path, 'template_registration.docx')):
        return True

    return False


def registration_form_build_in(course_details: str, debug: bool = False) -> dict:
    """
    Admission form that can be used inside the application
    :param course_details:
    :param debug: if we want debug information
    :return: information in a dictionary format from the form
    """
    local_dict = {}
    # TODO verify data integrate
    # TODO adapt date to the wanted format
    # TODO protect against empty or false data
    local_dict['Registration date:'] = str(datetime.date.today().day) + "/" + \
                                       str(datetime.date.today().month) + "/" + \
                                       str(datetime.date.today().year)
    local_dict['Name'] = input("Name: ")
    local_dict['Surname'] = input("Surname: ")
    local_dict['ID'] = input("ID: ")
    local_dict['Telephone'] = input("Telephone ")
    local_dict['Email'] = input("Email: ")
    local_dict['Wished Courses'] = input("Available courses:\n" + course_details
                                         + "\nDesired courses(The list of courses in order of priority):  ")
    if debug:
        print("DEBUG: Registration date: ", local_dict, '\n', '_' * 20)

    return local_dict


def student_registration(course_details: str, debug: bool = False) -> bool:
    """
    Function that handles the build_in form for adding an admission
    :param course_details: text to be written as details of the courses you can choose
    :param debug: if you want debug information
    :return: True if everything was OK
    """
    if debug:
        print("DEBUG: Available courses: ", course_details, '\n', '_' * 20)

    # read existing admissions to store
    if os.path.exists(os.path.join(os.getcwd(), "student.json")):
        json_file = open("student.json", 'r')
        text_list = json.load(json_file)
        json_file.close()
    else:
        text_list = []
    # get data from the online form
    form_output = registration_form_build_in(course_details, debug)
    # add new data from online form to the pending list
    text_list.append(form_output)
    # add new information to the json file
    json_file = open("student.json", 'w')
    json.dump(text_list, json_file)
    json_file.close()

    return True


def get_date_from_word_form(file_path: str, debug: bool = False) -> dict:
    """
    Parses a string and retrieves needed information
    :param file_path: path to docs file
    :param debug: if you want debug information
    :return: Valid information in a dictionary format
    """
    # link the variable to the docs document
    file_content = Document(file_path)
    # get all the paragraphs
    file_content = file_content.paragraphs

    local_dict = {'Registration date': remove_starting_space(file_content[1].text.split(':')[-1]),
                  'Name': remove_starting_space(file_content[2].text.split(':')[-1]),
                  'Surname': remove_starting_space(file_content[3].text.split(':')[-1]),
                  'ID': remove_starting_space(file_content[4].text.split(':')[-1]),
                  'Telephone': remove_starting_space(file_content[5].text.split(':')[-1]),
                  'Email': remove_starting_space(file_content[6].text.split(':')[-1]),
                  'Wished Courses': remove_starting_space(file_content[-1].text.split(':')[-1])}
    # second paragraph contains the date of admission
    # 3th paragraph contains the name of the applicant
    # 4th paragraph contains the surname of the applicant
    # 5th paragraph contains the surname of the applicant
    # 6th paragraph contains the telephone number of the applicant
    # 7th paragraph contains the email number of the applicant
    # the last paragraph contains the desired courses

    if debug:
        print("DEBUG: Output data from word file: ", local_dict, '\n', '_' * 20)

    return local_dict


def get_form_result_files(folder_path: str) -> list:
    """
    Parses the folder for applications and forms a list of all the paths
    :param folder_path: str the folder path
    :return: list of paths to files
    """
    list_files = []

    for root, dirs, files in os.walk(folder_path):
        for name in files:
            list_files.append(os.path.join(root, name))

    return list_files


def registration_students_word(folder_path: str, debug: bool = False) -> bool:
    """
    Parses all the documents in the folder with admissions and populates the correspondence json file.
    :param folder_path: location of enrollment
    :param debug: if you want debug information
    :return: True if all was good.
    """
    # read existing admissions to store
    if os.path.exists(os.path.join(os.getcwd(), "student.json")):
        json_file = open("student.json", 'r')
        text_list = json.load(json_file)
        json_file.close()
    else:
        text_list = []

    word_forms = get_form_result_files(os.path.join(os.getcwd(), folder_path, "applicants"))

    if debug:
        print("DEBUG: List of application forms are:  ", word_forms, '\n', '_' * 20)

    for file in word_forms:
        form_result = get_date_from_word_form(file, debug)
        # add new data from online form to the pending list
        text_list.append(form_result)

    # add new information to the json file
    try:
        json_file = open("student.json", 'w')
        json.dump(text_list, json_file)
        json_file.close()
    except IOError as e:
        print("ERROR %s: File does not appear to exist.", e.filename)

    return True


if __name__ == "__main__":
    # we use main for local testing of module
    registration_form_build_in()
